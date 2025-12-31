import tempfile
import os
import logging
from uuid import uuid4
from typing import List, Union
import json
import pandas as pd
import numpy as np
from fastapi import UploadFile
import fitz  # PyMuPDF
import docx
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from tabulate import tabulate
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores.utils import filter_complex_metadata

# Logger setup
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.DEBUG)
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
console_handler.setFormatter(formatter)
logger.addHandler(console_handler)

def extract_text_from_pdf(path):
    try:
        with fitz.open(path) as doc:
            logger.debug("Reading PDF using fitz")
            return "\n".join(page.get_text() for page in doc)
    except Exception as e:
        logger.error(f"Error reading PDF: {e}")
        raise

def extract_text_from_docx(path):
    try:
        logger.debug("Reading DOCX file")
        doc = docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        raise

def prepare_documents(text, filename):
    logger.debug("Splitting text into documents")
    splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=500, chunk_overlap=100
    )
    docs = splitter.create_documents([text])
    return [Document(page_content=doc.page_content, metadata={"source": filename}) for doc in docs]

def convert_excel_to_markdown_with_clusters_inline(path: str) -> str:
    try:
        wb = load_workbook(path, data_only=True)
        markdown_output = ""

        for sheet_name in wb.sheetnames:
            markdown_output += f"# {sheet_name}\n\n"
            ws = wb[sheet_name]

            clusters = []
            current_cluster = []

            for row in ws.iter_rows():
                if any(cell.value for cell in row):
                    current_cluster.append([cell.value for cell in row])
                elif current_cluster:
                    clusters.append(current_cluster)
                    current_cluster = []

            if current_cluster:
                clusters.append(current_cluster)

            for idx, cluster in enumerate(clusters):
                markdown_table = tabulate(cluster, headers='firstrow', tablefmt='github', showindex=False)
                markdown_output += f"### Table {idx + 1}\n\n"
                markdown_output += markdown_table + "\n\n"

        return markdown_output
    except Exception as e:
        return f"Error converting Excel to Markdown: {e}"

def extract_excel_clusters_with_metadata(filepath: Union[str, os.PathLike]) -> List[Document]:
    documents = []
    xls = pd.ExcelFile(filepath)

    for sheet_name in xls.sheet_names:
        try:
            sheet_df = xls.parse(sheet_name, header=None)
            cluster_id = 0
            in_block = False
            current_block = []

            for idx, row in sheet_df.iterrows():
                if row.isnull().all():
                    if in_block and current_block:
                        cluster_df = pd.DataFrame(current_block).dropna(how='all', axis=1)
                        if not cluster_df.empty:
                            try:
                                if cluster_df.shape[0] > 1:
                                    cluster_df.columns = cluster_df.iloc[0]
                                    cluster_df = cluster_df[1:]
                                else:
                                    cluster_df.columns = [f"Column {i}" for i in range(cluster_df.shape[1])]

                                cluster_df = cluster_df.reset_index(drop=True)
                                metadata = {
                                    "source": os.path.basename(filepath),
                                    "sheet_name": sheet_name,
                                    "type": "table",
                                    "cluster_id": str(cluster_id),
                                    "columns": safe_column_list(cluster_df.columns),
                                    "shape": stringify_shape(cluster_df),
                                }

                                if is_valid_metadata(metadata):
                                    content = summarize_cluster(cluster_df)
                                    documents.append(Document(page_content=cluster_df.to_markdown(index=False), metadata=metadata))
                                cluster_id += 1
                            except Exception as e:
                                logger.warning(f"Error processing cluster in sheet {sheet_name}: {e}")
                        current_block = []
                        in_block = False
                else:
                    current_block.append(row)
                    in_block = True

        except Exception as e:
            logger.error(f"Failed to process sheet {sheet_name}: {e}")

    return documents

def safe_column_list(columns) -> List[str]:
    return [str(col).strip() if col is not None and str(col).strip().lower() != 'nan' else '' for col in columns]

def stringify_shape(df: pd.DataFrame) -> str:
    return f"{df.shape[0]} rows × {df.shape[1]} cols"

def is_valid_metadata(metadata: dict) -> bool:
    try:
        json.dumps(metadata)
        return True
    except Exception as e:
        logger.warning(f"Invalid metadata: {metadata} ({type(metadata)}) -> {e}")
        return False

def summarize_cluster(df: pd.DataFrame) -> str:
    sample = df.head(2).to_string(index=False)
    return f"Table with {df.shape[0]} rows and {df.shape[1]} columns.\nSample:\n{sample}"

# processor.py
async def process_file(file: UploadFile) -> str:
    """Process a single file and return markdown string"""
    logger.info(f"Processing file: {file.filename}")
    suffix = file.filename.split(".")[-1].lower()

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
        file_bytes = await file.read()
        tmp.write(file_bytes)
        tmp.flush()
        os.fsync(tmp.fileno())
        path = tmp.name

    try:
        if suffix == "pdf":
            try:
                loader = PyPDFLoader(path)
                docs = loader.load()
                markdown_data = docs
            except Exception as e:
                logger.warning(f"PyPDFLoader failed: {e}, falling back to fitz")
                text = extract_text_from_pdf(path)
                docs = prepare_documents(text, file.filename)
                markdown_data = docs
        elif suffix in ["docx", "doc"]:
            text = extract_text_from_docx(path)
            docs = prepare_documents(text, file.filename)
            markdown_data = docs
        elif suffix in ["xlsx", "xls"]:
            try:
                docs = extract_excel_clusters_with_metadata(path)
                markdown_data = convert_excel_to_markdown_with_clusters_inline(path)
            finally:
                # Make sure Excel file is closed
                import gc
                gc.collect()
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        return markdown_data

    finally:
        try:
            os.remove(path)
            logger.debug(f"Deleted temporary file: {path}")
        except PermissionError:
            logger.warning(f"Could not immediately delete file: {path}, will be cleaned up later")
            # Schedule file for deletion on next restart
            try:
                import atexit
                atexit.register(lambda p=path: os.remove(p) if os.path.exists(p) else None)
            except Exception as e:
                logger.error(f"Error scheduling file cleanup: {e}")
    """Process a single file and return markdown string"""
    logger.info(f"Processing file: {file.filename}")
    suffix = file.filename.split(".")[-1].lower()

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
        file_bytes = await file.read()
        tmp.write(file_bytes)
        tmp.flush()
        os.fsync(tmp.fileno())
        path = tmp.name
        logger.debug(f"Saved temporary file to: {path}")

    try:
        if suffix == "pdf":
            try:
                loader = PyPDFLoader(path)
                docs = loader.load()
                markdown_data = docs
            except Exception as e:
                logger.warning(f"PyPDFLoader failed: {e}, falling back to fitz")
                text = extract_text_from_pdf(path)
                docs = prepare_documents(text, file.filename)
                markdown_data = docs
        elif suffix in ["docx", "doc"]:
            text = extract_text_from_docx(path)
            docs = prepare_documents(text, file.filename)
            markdown_data = docs
        elif suffix in ["xlsx", "xls"]:
            docs = extract_excel_clusters_with_metadata(path)
            markdown_data = convert_excel_to_markdown_with_clusters_inline(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        return markdown_data

    finally:
        os.remove(path)
        logger.debug(f"Deleted temporary file: {path}")